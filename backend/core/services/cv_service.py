import json
import mimetypes
import os
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from werkzeug.datastructures import FileStorage

from ..security import slugify

ALLOWED_EXTENSIONS = {"pdf", "doc", "docx"}


def allowed_resume_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def ensure_upload_dir(base_dir: str) -> Path:
    path = Path(base_dir)
    path.mkdir(parents=True, exist_ok=True)
    return path


def save_uploaded_file(file_storage: FileStorage, base_dir: str, prefix: str) -> tuple[str, str, str]:
    safe_name = slugify(Path(file_storage.filename).stem)
    ext = Path(file_storage.filename).suffix.lower()
    file_name = f"{prefix}-{safe_name}{ext}"
    upload_dir = ensure_upload_dir(base_dir)
    destination = upload_dir / file_name
    file_storage.save(destination)
    mime_type = mimetypes.guess_type(str(destination))[0] or "application/octet-stream"
    return file_name, str(destination), mime_type


def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs]).strip()


def extract_text_from_doc(path: str) -> str:
    # Best-effort support for .doc uploads.
    try:
        import subprocess

        output_dir = Path(path).parent
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                path,
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        converted = Path(path).with_suffix(".docx")
        if converted.exists():
            return extract_text_from_docx(str(converted))
    except Exception:
        pass
    return ""


def extract_text_from_upload(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext == ".docx":
        return extract_text_from_docx(path)
    if ext == ".doc":
        return extract_text_from_doc(path)
    return ""


def generate_pdf_from_resume(data: dict, output_path: str) -> str:
    def _resolve_pdf_font() -> str:
        candidates = [
            os.getenv("PDF_FONT_PATH", "").strip(),
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/tahoma.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for raw in candidates:
            if not raw:
                continue
            path = Path(raw)
            if not path.exists():
                continue
            font_name = f"ResumeFont-{path.stem}"
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
        return "Helvetica"

    def _text(value):
        if value is None:
            return ""
        if isinstance(value, list):
            return ", ".join(str(item) for item in value)
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False)
        return str(value)

    def _wrap_text(text: str, font_name: str, font_size: int, max_width: float) -> list[str]:
        normalized = str(text or "").replace("\r\n", "\n")
        output: list[str] = []
        for raw_line in normalized.split("\n"):
            line = raw_line.strip()
            if not line:
                output.append("")
                continue

            current = ""
            for word in line.split():
                candidate = word if not current else f"{current} {word}"
                if pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
                    current = candidate
                    continue

                if current:
                    output.append(current)

                if pdfmetrics.stringWidth(word, font_name, font_size) <= max_width:
                    current = word
                    continue

                chunk = ""
                for char in word:
                    maybe = f"{chunk}{char}"
                    if pdfmetrics.stringWidth(maybe, font_name, font_size) <= max_width:
                        chunk = maybe
                    else:
                        if chunk:
                            output.append(chunk)
                        chunk = char
                current = chunk

            if current:
                output.append(current)
        return output

    font_name = _resolve_pdf_font()
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    margin_x = 36
    margin_top = 36
    margin_bottom = 36
    page_w = width - (margin_x * 2)
    y = height - margin_top

    def _draw_wrapped(text: str, x: float, y_pos: float, max_width: float, size: int = 11, leading: int = 15):
        c.setFont(font_name, size)
        lines = _wrap_text(text, font_name, size, max_width)
        for line in lines:
            c.drawString(x, y_pos, line)
            y_pos -= leading
        return y_pos

    def _safe(value):
        return _text(value).strip() or "Chua cap nhat"

    template = data.get("template") or {}
    template_name = _safe(template.get("name") or data.get("template_name") or "CV Template")
    template_slug = str(template.get("slug") or data.get("template_slug") or template_name).strip().lower().replace(" ", "-")
    headline = _safe(data.get("headline"))
    full_name = _safe(data.get("full_name") or "Resume")

    def _section_title(title: str, x: float, y_pos: float, color_hex: str = "#1e3a8a"):
        c.setFillColor(colors.HexColor(color_hex))
        c.setFont(font_name, 13)
        c.drawString(x, y_pos, title)
        c.setFillColor(colors.black)
        return y_pos - 18

    def _draw_modern_layout(theme: dict[str, str]):
        nonlocal y
        c.setFillColor(colors.HexColor(theme["band_bg"]))
        c.rect(margin_x, y - 22, page_w, 22, fill=1, stroke=0)
        c.setFillColor(colors.HexColor(theme["band_text"]))
        c.setFont(font_name, 11)
        c.drawString(margin_x + 10, y - 14, template_name.upper())
        c.drawRightString(margin_x + page_w - 10, y - 14, (data.get("current_title") or headline).upper())
        y -= 34

        c.setFillColor(colors.HexColor(theme["avatar"]))
        c.circle(margin_x + 28, y - 28, 18, stroke=0, fill=1)
        c.setFillColor(colors.black)
        c.setFont(font_name, 24)
        c.drawString(margin_x + 58, y - 18, full_name)
        c.setFillColor(colors.HexColor(theme["accent"]))
        c.setFont(font_name, 12)
        c.drawString(margin_x + 58, y - 38, headline)
        c.setFillColor(colors.black)

        contact_lines = [
            f"Ngay sinh: {_safe(data.get('dob'))}",
            f"Gioi tinh: {_safe(data.get('gender'))}",
            f"So dien thoai: {_safe(data.get('phone'))}",
            f"Email: {_safe(data.get('email'))}",
            f"Dia chi: {_safe(data.get('address'))}",
        ]
        line_y = y - 58
        for line in contact_lines:
            c.setFont(font_name, 11)
            c.drawString(margin_x + 58, line_y, line)
            line_y -= 14

        pill_y = line_y - 2
        pills = [
            _safe(data.get("expected_salary")),
            _safe(data.get("desired_location")),
            f"{int(data.get('years_experience') or 0)} nam kinh nghiem",
        ]
        x_cursor = margin_x + 58
        c.setFont(font_name, 10)
        for pill in pills:
            w = pdfmetrics.stringWidth(pill, font_name, 10) + 14
            c.setFillColor(colors.HexColor(theme["pill_bg"]))
            c.roundRect(x_cursor, pill_y - 10, w, 16, 8, fill=1, stroke=0)
            c.setFillColor(colors.HexColor(theme["accent"]))
            c.drawString(x_cursor + 7, pill_y - 4, pill)
            x_cursor += w + 8
        c.setFillColor(colors.black)

        y = pill_y - 20
        c.setStrokeColor(colors.HexColor("#d1d5db"))
        c.line(margin_x, y, margin_x + page_w, y)
        y -= 10

        left_w = page_w * 0.38
        gap = 16
        right_w = page_w - left_w - gap
        left_x = margin_x
        right_x = margin_x + left_w + gap
        left_y = y
        right_y = y

        c.setFillColor(colors.HexColor(theme["left_bg"]))
        c.rect(left_x, margin_bottom, left_w, y - margin_bottom + 6, fill=1, stroke=0)
        c.setFillColor(colors.black)

        left_y = _section_title("Muc tieu nghe nghiep", left_x + 12, left_y, theme["accent_dark"])
        left_y = _draw_wrapped(_safe(data.get("summary")), left_x + 12, left_y, left_w - 24)
        left_y -= 8
        left_y = _section_title("Ky nang", left_x + 12, left_y, theme["accent_dark"])
        left_y = _draw_wrapped(_safe(data.get("skills")), left_x + 12, left_y, left_w - 24)

        right_y = _section_title("Kinh nghiem lam viec", right_x, right_y, theme["accent_dark"])
        right_y = _draw_wrapped(_safe(data.get("experience")), right_x, right_y, right_w)
        right_y -= 8
        right_y = _section_title("Hoc van", right_x, right_y, theme["accent_dark"])
        _draw_wrapped(_safe(data.get("education")), right_x, right_y, right_w)

    def _draw_ats_layout():
        nonlocal y
        c.setStrokeColor(colors.HexColor("#111827"))
        c.setLineWidth(1)
        c.setFont(font_name, 26)
        c.drawString(margin_x, y, full_name)
        y -= 20
        c.setFont(font_name, 12)
        c.drawString(margin_x, y, headline)
        y -= 12
        c.line(margin_x, y, margin_x + page_w, y)
        y -= 16

        c.setFont(font_name, 11)
        meta = [
            f"Email: {_safe(data.get('email'))}",
            f"Phone: {_safe(data.get('phone'))}",
            f"Address: {_safe(data.get('address'))}",
            f"Current title: {_safe(data.get('current_title'))}",
            f"Experience: {int(data.get('years_experience') or 0)} nam",
        ]
        for item in meta:
            c.drawString(margin_x, y, item)
            y -= 14
        y -= 6

        for title, body in [
            ("SUMMARY", _safe(data.get("summary"))),
            ("SKILLS", _safe(data.get("skills"))),
            ("EXPERIENCE", _safe(data.get("experience"))),
            ("EDUCATION", _safe(data.get("education"))),
        ]:
            y = _section_title(title, margin_x, y, "#111827")
            y = _draw_wrapped(body, margin_x, y, page_w)
            y -= 8

    def _draw_creative_layout():
        nonlocal y
        c.setFillColor(colors.HexColor("#111827"))
        c.rect(margin_x, y - 48, page_w, 48, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont(font_name, 22)
        c.drawString(margin_x + 12, y - 18, full_name)
        c.setFont(font_name, 11)
        c.drawString(margin_x + 12, y - 36, headline)
        c.setFillColor(colors.black)
        y -= 62

        left_w = page_w * 0.62
        gap = 14
        right_w = page_w - left_w - gap
        left_x = margin_x
        right_x = margin_x + left_w + gap
        left_y = y
        right_y = y

        left_y = _section_title("Profile", left_x, left_y, "#111827")
        left_y = _draw_wrapped(_safe(data.get("summary")), left_x, left_y, left_w)
        left_y -= 8
        left_y = _section_title("Experience", left_x, left_y, "#111827")
        left_y = _draw_wrapped(_safe(data.get("experience")), left_x, left_y, left_w)

        right_y = _section_title("Contact", right_x, right_y, "#111827")
        right_y = _draw_wrapped(
            f"Email: {_safe(data.get('email'))}\nPhone: {_safe(data.get('phone'))}\nAddress: {_safe(data.get('address'))}",
            right_x,
            right_y,
            right_w,
        )
        right_y -= 8
        right_y = _section_title("Skills", right_x, right_y, "#111827")
        right_y = _draw_wrapped(_safe(data.get("skills")), right_x, right_y, right_w)
        right_y -= 8
        right_y = _section_title("Education", right_x, right_y, "#111827")
        _draw_wrapped(_safe(data.get("education")), right_x, right_y, right_w)

    modern_themes = {
        "modern-blue": {
            "band_bg": "#dbeafe",
            "band_text": "#1e3a8a",
            "avatar": "#2563eb",
            "accent": "#1d4ed8",
            "accent_dark": "#1e3a8a",
            "left_bg": "#f8fbff",
            "pill_bg": "#eff6ff",
        },
        "product-designer": {
            "band_bg": "#e2e8f0",
            "band_text": "#0f172a",
            "avatar": "#334155",
            "accent": "#0f172a",
            "accent_dark": "#0f172a",
            "left_bg": "#f8fafc",
            "pill_bg": "#f1f5f9",
        },
        "data-analyst": {
            "band_bg": "#bfdbfe",
            "band_text": "#1e3a8a",
            "avatar": "#1e40af",
            "accent": "#1d4ed8",
            "accent_dark": "#1e3a8a",
            "left_bg": "#eff6ff",
            "pill_bg": "#dbeafe",
        },
        "hr-executive": {
            "band_bg": "#fecdd3",
            "band_text": "#9f1239",
            "avatar": "#be123c",
            "accent": "#be123c",
            "accent_dark": "#9f1239",
            "left_bg": "#fff1f2",
            "pill_bg": "#ffe4e6",
        },
        "marketing-pro": {
            "band_bg": "#fde68a",
            "band_text": "#92400e",
            "avatar": "#f97316",
            "accent": "#c2410c",
            "accent_dark": "#92400e",
            "left_bg": "#fffbeb",
            "pill_bg": "#fef3c7",
        },
    }

    if "ats" in template_slug:
        _draw_ats_layout()
    elif "creative" in template_slug or "minimal" in template_slug:
        _draw_creative_layout()
    else:
        _draw_modern_layout(modern_themes.get(template_slug, modern_themes["modern-blue"]))

    c.save()
    return output_path


def generate_docx_from_resume(data: dict, output_path: str) -> str:
    def _text(value):
        if value is None:
            return ""
        if isinstance(value, list):
            return ", ".join(str(item) for item in value)
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False)
        return str(value)

    doc = Document()
    doc.add_heading(_text(data.get("full_name", "Resume")) or "Resume", level=1)
    if data.get("headline"):
        doc.add_paragraph(_text(data["headline"]))
    for label, key in [
        ("Summary", "summary"),
        ("Skills", "skills"),
        ("Experience", "experience"),
        ("Education", "education"),
    ]:
        doc.add_heading(label, level=2)
        doc.add_paragraph(_text(data.get(key, "")))
    doc.save(output_path)
    return output_path

